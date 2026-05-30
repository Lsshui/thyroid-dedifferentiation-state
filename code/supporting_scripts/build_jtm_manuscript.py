from pathlib import Path
import csv
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parents[2]
ASSET = BASE / "analysis" / "manuscript_assets"
OUTDIR = ASSET / "draft_manuscript"
OUTDIR.mkdir(parents=True, exist_ok=True)

MANUSCRIPT_DOCX = OUTDIR / "JTM_dedifferentiation_state_thyroid_cancer_revised_v2.docx"
NOTES_DOCX = OUTDIR / "JTM_dedifferentiation_state_revision_notes_v2.docx"
MANUSCRIPT_MD = OUTDIR / "JTM_dedifferentiation_state_thyroid_cancer_revised_v2.md"


def pval(p):
    try:
        x = float(p)
    except Exception:
        return str(p)
    if x < 1e-4:
        return f"{x:.2e}"
    return f"{x:.3f}"


def pct(x):
    try:
        return f"{float(x) * 100:.0f}%"
    except Exception:
        return str(x)


def read_csv(path):
    with open(path, newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


table1_raw = read_csv(ASSET / "tables" / "Table_1_cohorts.csv")
table2_raw = read_csv(ASSET / "tables" / "Table_2_key_evidence.csv")
table3_raw = read_csv(ASSET / "tables" / "Table_3_target_prioritization.csv")

def fmt_int(x):
    try:
        return f"{int(float(x)):,}"
    except Exception:
        return str(x)


table1 = [
    [r["Cohort"], r["Modality"], r["Role"], fmt_int(r["N samples/cells"])]
    for r in table1_raw
]

table2 = [
    [
        "RAIR residual -> iodine-handling",
        "TCGA-THCA",
        "beta",
        "-1.016 (-1.160 to -0.872)",
        "4.46e-37",
        "Composite score retains iodine-handling loss beyond MAPK, CAF, immune and driver status.",
    ],
    [
        "RAIR residual -> N1",
        "TCGA-THCA",
        "OR",
        "2.158 (1.345 to 3.461)",
        "0.001",
        "Residual composite score remains associated with nodal disease.",
    ],
    [
        "ATC vs PTC",
        "GSE33630",
        "AUC",
        "0.885",
        "1.69e-05",
        "External validation of the dedifferentiation-state score.",
    ],
    [
        "ATC vs PDTC",
        "GSE76039",
        "AUC",
        "0.871",
        "4.70e-05",
        "External validation in a high-grade thyroid cancer contrast.",
    ],
    [
        "copyKAT aneuploid vs diploid epithelial cells",
        "GSE232237",
        "Median difference",
        "3.079",
        "4.18e-11",
        "Composite score is enriched in copyKAT-aneuploid epithelial/thyroid cells.",
    ],
]

target_order = ["SLC7A5", "SLC1A5", "MTHFD2", "AXL", "FN1", "FOSL1", "ITGB1"]
table3_by_gene = {r["Gene"]: r for r in table3_raw}
hpa_rows = read_csv(BASE / "analysis" / "results" / "HPA_target_thyroid_cancer_pathology.csv")
hpa_by_gene = {r["gene"]: r for r in hpa_rows}
axis_short = {
    "SLC7A5": "Amino-acid uptake",
    "SLC1A5": "Amino-acid uptake",
    "MTHFD2": "Mitochondrial one-carbon",
    "AXL": "RTK/EMT survival",
    "FN1": "ECM/state",
    "FOSL1": "AP-1/MAPK state",
    "ITGB1": "ECM adhesion",
}
table3 = []
for g in target_order:
    r = table3_by_gene[g]
    hpa = hpa_by_gene.get(g, {})
    high_med = int(float(hpa.get("pathology_high", 0))) + int(float(hpa.get("pathology_medium", 0)))
    hpa_total = int(float(hpa.get("pathology_total", 0))) if hpa else 0
    hpa_text = f"{high_med}/{hpa_total}" if hpa_total else "NA"
    chembl = r.get("Actionability bin", "")
    records = r.get("ChEMBL activity records", "")
    if records:
        chembl = f"{chembl}; {records} records"
    interp = {
        "SLC7A5": "Strongest integrated transporter candidate; protein and chemistry evidence support prioritization.",
        "SLC1A5": "Consistent transporter signal with preclinical chemical tractability; supportive metabolic node.",
        "MTHFD2": "Transcriptomic metabolic candidate; HPA and ChEMBL evidence remain sparse.",
        "AXL": "Clinically druggable RTK/EMT node; thyroid cancer HPA staining was not supportive in the small series.",
        "FN1": "Robust ecology/state anchor rather than a single-agent drug target.",
        "FOSL1": "AP-1/MAPK state marker; best used as a transcriptional readout.",
        "ITGB1": "Adhesion/ecology marker; not prioritized as a direct monotherapy target.",
    }[g]
    table3.append([
        g,
        "Core translational" if "1A" in r["Tier"] else "Ecology/state",
        axis_short[g],
        f"{float(r['TCGA rho RAIR-like']):.3f}",
        f"{r['External validations']}/2",
        hpa_text,
        chembl if chembl else "Not prioritized",
        interp,
    ])


references = [
    ("Amelio I", "2014", "Serine and glycine metabolism in cancer.", "Trends in Biochemical Sciences", "10.1016/j.tibs.2014.02.004", "24657017"),
    ("Antony J", "2017", "AXL-Driven EMT State as a Targetable Conduit in Cancer.", "Cancer Research", "10.1158/0008-5472.CAN-17-0392", "28667075"),
    ("Avilla E", "2011", "Activation of TYRO3/AXL tyrosine kinase receptors in thyroid cancer.", "Cancer Research", "10.1158/0008-5472.CAN-10-2186", "21343401"),
    ("Bible KC", "2021", "2021 American Thyroid Association Guidelines for Management of Patients with Anaplastic Thyroid Cancer.", "Thyroid", "10.1089/thy.2020.0944", "33728999"),
    ("Bhalla S", "2023", "Phase 1 trial of bemcentinib (BGB324), a first-in-class, selective AXL inhibitor, with docetaxel in patients with previously treated advanced non-small cell lung cancer.", "Lung Cancer", "10.1016/j.lungcan.2023.107291", "37423058"),
    ("Brose MS", "2014", "Sorafenib in radioactive iodine-refractory, locally advanced or metastatic differentiated thyroid cancer: a randomised, double-blind, phase 3 trial.", "Lancet", "10.1016/S0140-6736(14)60421-9", "24768112"),
    ("Cabanillas ME", "2016", "Thyroid cancer.", "Lancet", "10.1016/S0140-6736(16)30172-6", "27240885"),
    ("Cancer Genome Atlas Research Network", "2014", "Integrated genomic characterization of papillary thyroid carcinoma.", "Cell", "10.1016/j.cell.2014.09.050", "25417114"),
    ("Chakravarty D", "2011", "Small-molecule MAPK inhibitors restore radioiodine incorporation in mouse thyroid cancers with conditional BRAF activation.", "Journal of Clinical Investigation", "10.1172/JCI46382", "22105174"),
    ("Colaprico A", "2016", "TCGAbiolinks: an R/Bioconductor package for integrative analysis of TCGA data.", "Nucleic Acids Research", "10.1093/nar/gkv1507", "26704973"),
    ("Colombo C", "2020", "The molecular and gene/miRNA expression profiles of radioiodine resistant papillary thyroid cancer.", "Journal of Experimental & Clinical Cancer Research", "10.1186/s13046-020-01757-x", "33198784"),
    ("Davis S", "2007", "GEOquery: a bridge between the Gene Expression Omnibus and Bioconductor.", "Bioinformatics", "10.1093/bioinformatics/btm254", "17496320"),
    ("Ducker GS", "2017", "One-Carbon Metabolism in Health and Disease.", "Cell Metabolism", "10.1016/j.cmet.2016.08.009", "27641100"),
    ("Durante C", "2006", "Long-term outcome of 444 patients with distant metastases from papillary and follicular thyroid carcinoma: benefits and limits of radioiodine therapy.", "Journal of Clinical Endocrinology and Metabolism", "10.1210/jc.2005-2838", "16684830"),
    ("Fagin JA", "2016", "Biologic and Clinical Perspectives on Thyroid Cancer.", "New England Journal of Medicine", "10.1056/NEJMra1501993", "27626519"),
    ("Gao R", "2021", "Delineating copy number and clonal substructure in human tumors from single-cell transcriptomes.", "Nature Biotechnology", "10.1038/s41587-020-00795-2", "33462507"),
    ("Gaulton A", "2012", "ChEMBL: a large-scale bioactivity database for drug discovery.", "Nucleic Acids Research", "10.1093/nar/gkr777", "21948594"),
    ("Gjerdrum C", "2010", "Axl is an essential epithelial-to-mesenchymal transition-induced regulator of breast cancer metastasis and patient survival.", "Proceedings of the National Academy of Sciences of the United States of America", "10.1073/pnas.0909333107", "20080645"),
    ("Haugen BR", "2016", "2015 American Thyroid Association Management Guidelines for Adult Patients with Thyroid Nodules and Differentiated Thyroid Cancer.", "Thyroid", "10.1089/thy.2015.0020", "26462967"),
    ("Hebrant A", "2012", "mRNA expression in papillary and anaplastic thyroid carcinoma: molecular anatomy of a killing switch.", "PLoS One", "10.1371/journal.pone.0037807", "23115614"),
    ("Ho AL", "2013", "Selumetinib-enhanced radioiodine uptake in advanced thyroid cancer.", "New England Journal of Medicine", "10.1056/NEJMoa1209288", "23406027"),
    ("Jin Y", "2018", "Radioiodine refractory differentiated thyroid cancer.", "Critical Reviews in Oncology/Hematology", "10.1016/j.critrevonc.2018.03.012", "29650270"),
    ("Kanai Y", "2022", "Amino acid transporter LAT1 (SLC7A5) as a molecular target for cancer diagnosis and therapeutics.", "Pharmacology & Therapeutics", "10.1016/j.pharmthera.2021.107964", "34390745"),
    ("Kawai J", "2019", "Discovery of a potent, selective, and orally available MTHFD2 inhibitor (DS18561882) with in vivo antitumor activity.", "Journal of Medicinal Chemistry", "10.1021/acs.jmedchem.9b01113", "31638799"),
    ("Landa I", "2016", "Genomic and transcriptomic hallmarks of poorly differentiated and anaplastic thyroid cancers.", "Journal of Clinical Investigation", "10.1172/JCI85271", "26878173"),
    ("Leboulleux S", "2023", "A Phase II Redifferentiation Trial with Dabrafenib-Trametinib and 131I in Metastatic Radioactive Iodine Refractory BRAF p.V600E-Mutated Differentiated Thyroid Cancer.", "Clinical Cancer Research", "10.1158/1078-0432.CCR-23-0046", "37074727"),
    ("Liberzon A", "2015", "The Molecular Signatures Database Hallmark gene set collection.", "Cell Systems", "10.1016/j.cels.2015.12.004", "26771021"),
    ("Locasale JW", "2013", "Serine, glycine and one-carbon units: cancer metabolism in full circle.", "Nature Reviews Cancer", "10.1038/nrc3557", "23822983"),
    ("Lu L", "2023", "Anaplastic transformation in thyroid cancer revealed by single-cell transcriptomics.", "Journal of Clinical Investigation", "10.1172/JCI169653", "37053016"),
    ("Luo H", "2021", "Characterizing dedifferentiation of thyroid cancer by integrated analysis.", "Science Advances", "10.1126/sciadv.abf3657", "34321197"),
    ("Nachef M", "2021", "Targeting SLC1A5 and SLC3A2/SLC7A5 as a potential strategy to strengthen anti-tumor immunity in the tumor microenvironment.", "Frontiers in Immunology", "10.3389/fimmu.2021.624324", "33953707"),
    ("Nilsson R", "2014", "Metabolic enzyme expression highlights a key role for MTHFD2 and the mitochondrial folate pathway in cancer.", "Nature Communications", "10.1038/ncomms4128", "24451681"),
    ("Okano N", "2020", "Biomarker analyses in patients with advanced solid tumors treated with the LAT1 inhibitor JPH203.", "In Vivo", "10.21873/invivo.12077", "32871789"),
    ("Pu W", "2021", "Single-cell transcriptomic analysis of the tumor ecosystems underlying initiation and progression of papillary thyroid carcinoma.", "Nature Communications", "10.1038/s41467-021-26343-3", "34663816"),
    ("Ramos L", "2024", "Targeting MTHFD2 to exploit cancer-specific metabolism and the DNA damage response.", "Cancer Research", "10.1158/0008-5472.CAN-23-1290", "37922465"),
    ("Ritchie ME", "2015", "limma powers differential expression analyses for RNA-sequencing and microarray studies.", "Nucleic Acids Research", "10.1093/nar/gkv007", "25605792"),
    ("Scalise M", "2018", "The Human SLC7A5 (LAT1): The intriguing histidine/large neutral amino acid transporter and its relevance to human health.", "Frontiers in Chemistry", "10.3389/fchem.2018.00243", "29988369"),
    ("Schlumberger M", "2015", "Lenvatinib versus placebo in radioiodine-refractory thyroid cancer.", "New England Journal of Medicine", "10.1056/NEJMoa1406470", "25671254"),
    ("Subbiah V", "2018", "Dabrafenib and trametinib treatment in patients with locally advanced or metastatic BRAF V600-mutant anaplastic thyroid cancer.", "Journal of Clinical Oncology", "10.1200/JCO.2017.73.6785", "29072975"),
    ("Subbiah V", "2022", "Dabrafenib plus trametinib in patients with BRAF V600E-mutant anaplastic thyroid cancer: updated analysis from the phase II ROAR basket study.", "Annals of Oncology", "10.1016/j.annonc.2021.12.014", "35026411"),
    ("Subramanian A", "2005", "Gene set enrichment analysis: a knowledge-based approach for interpreting genome-wide expression profiles.", "Proceedings of the National Academy of Sciences of the United States of America", "10.1073/pnas.0506580102", "16199517"),
    ("Tchekmedyian V", "2022", "Enhancing radioiodine incorporation in BRAF-mutant, radioiodine-refractory thyroid cancers with vemurafenib and the anti-ErbB3 monoclonal antibody CDX-3379: results of a pilot clinical trial.", "Thyroid", "10.1089/thy.2021.0565", "35045748"),
    ("Uhlen M", "2015", "Proteomics. Tissue-based map of the human proteome.", "Science", "10.1126/science.1260419", "25613900"),
    ("Uhlen M", "2017", "A pathology atlas of the human cancer transcriptome.", "Science", "10.1126/science.aan2507", "28818916"),
    ("Wang Q", "2015", "L-type amino acid transport and cancer: targeting the mTORC1 pathway to inhibit neoplasia.", "American Journal of Cancer Research", "", "26101697"),
    ("Zdrazil B", "2024", "The ChEMBL Database in 2023: a drug discovery platform spanning multiple bioactivity data types and time periods.", "Nucleic Acids Research", "10.1093/nar/gkad1004", "37933841"),
    ("Zhu Z", "2020", "More than a metabolic enzyme: MTHFD2 as a novel target for anticancer therapy?", "Frontiers in Oncology", "10.3389/fonc.2020.00658", "32411609"),
]


title = "Integrated transcriptomic and single-cell mapping of a one-carbon-metabolism-enriched dedifferentiation state in thyroid cancer prioritizes SLC7A5, MTHFD2 and AXL"

authors = (
    "Shi Yin1,2,3*, Li Shuishi4, Chang Shi2,4*"
)
affiliations = [
    "1 Department of Pharmacy, Xiangya Hospital Central South University, Changsha, Hunan, China.",
    "2 National Clinical Research Center for Geriatric Disorders, Xiangya Hospital, Central South University, Changsha, Hunan, China.",
    "3 The Hunan Institute of Pharmacy Practice and Clinical Research, Changsha, Hunan, China.",
    "4 Department of General Surgery, Xiangya Hospital Central South University, Changsha, Hunan, China.",
    "* Correspondence: Chang Shi, changshi@csu.edu.cn; Shi Yin, shiyin910515@csu.edu.cn.",
    "ORCID: Chang Shi, 0000-0001-5291-0151; Shi Yin, 0000-0002-5763-3488.",
]


abstract_sections = {
    "Background": "Loss of thyroid differentiation underlies aggressive and radioiodine-refractory thyroid cancer. Public transcriptomic cohorts with direct radioiodine-response labels remain small. We asked whether a composite dedifferentiation-state score, integrating one-carbon metabolism, MAPK activity, EMT and loss of thyroid differentiation, could map the PTC-PDTC-ATC continuum, localize to malignant epithelial compartments and prioritize translationally tractable nodes.",
    "Methods": "We integrated TCGA-THCA bulk RNA-seq, three GEO bulk cohorts (GSE151179, GSE33630 and GSE76039), two single-cell RNA-seq cohorts (GSE184362 and GSE232237), copyKAT-based malignant-cell inference, Human Protein Atlas immunohistochemistry and ChEMBL target annotation. The composite score, retained as a RAIR-like shorthand, was defined as z(One-carbon metabolism) + z(MAPK) + z(dedifferentiation/EMT) - z(thyroid differentiation score). Associations were tested using rank-based statistics, logistic regression, pathway enrichment and composition-adjusted single-cell analyses.",
    "Results": "In TCGA-THCA, the composite score correlated inversely with iodine-handling score (rho=-0.341, P=4.44e-15) and remained associated with iodine-handling loss after adjustment for CAF, immune checkpoint, BRAF and RAS status (beta=-0.192, 95% CI -0.282 to -0.102, P=3.03e-05). A MAPK-adjusted residual retained a strong association with iodine-handling score (beta=-1.016, 95% CI -1.160 to -0.872, P=4.46e-37) and nodal disease (OR=2.158, 95% CI 1.345 to 3.461, P=0.001). GSE151179 contained only four RAI-avid tumors and had limited power for RAI-response biomarker testing; neither TDS nor the composite score separated RAI-refractory from RAI-avid tumors. In contrast, the same score separated ATC from PTC in GSE33630 (AUC=0.885, P=1.69e-05) and ATC from PDTC in GSE76039 (AUC=0.871, P=4.70e-05). Single-cell analyses localized high-score cells mainly to epithelial/thyroid compartments, with ecosystem contributions. In GSE232237, copyKAT-aneuploid epithelial cells showed higher scores than diploid epithelial cells (median difference=3.079, P=4.18e-11). Integrated annotation prioritized SLC7A5, SLC1A5, MTHFD2 and AXL, with FN1, FOSL1 and ITGB1 serving mainly as ecology or state anchors.",
    "Conclusions": "A composite dedifferentiation-state score reproducibly maps thyroid cancer dedifferentiation across bulk and single-cell cohorts and prioritizes SLC7A5, MTHFD2, SLC1A5 and AXL as translational nodes warranting functional validation. Although the score did not classify radioiodine refractoriness in the only available public RAI-response cohort, it provides a reproducible framework for stratifying dedifferentiation biology and for designing future RAI-response studies."
}

keywords = "Thyroid cancer; radioiodine-refractory; dedifferentiation; one-carbon metabolism; SLC7A5; MTHFD2; AXL; copyKAT; single-cell RNA-seq"


sections = [
    ("Introduction", [
        "Differentiated thyroid cancer is usually indolent, yet a clinically important subset loses radioiodine avidity or progresses despite radioiodine-based therapy. The therapeutic value and limits of radioiodine in metastatic differentiated thyroid cancer were already clear in long-term clinical series, where outcome depended strongly on the persistence of iodine uptake {Durante, 2006}. Contemporary guidelines still define loss of radioiodine uptake and progression after radioiodine as central features of radioiodine-refractory disease {Haugen, 2016}. Once tumors enter this state, treatment shifts toward systemic kinase inhibition or clinical trials, with sorafenib and lenvatinib improving progression-free survival but not resolving the biologic heterogeneity of the disease {Brose, 2014; Schlumberger, 2015}.",
        "Molecular studies have placed thyroid differentiation loss at the center of this problem. TCGA-THCA separated papillary thyroid carcinoma into BRAF-like and RAS-like transcriptional states and introduced the thyroid differentiation score as a quantitative readout of iodine-handling and thyroid-lineage programs {Cancer Genome Atlas Research Network, 2014}. Poorly differentiated and anaplastic thyroid cancers extend this continuum with stronger cell-cycle, MAPK, PI3K, TP53 and chromatin-associated alterations, together with marked loss of thyroid-lineage expression {Landa, 2016}. Integrated studies of thyroid cancer dedifferentiation further support the idea that progression is not a simple histologic jump, but a graded transcriptional and genomic transition {Luo, 2021}.",
        "The clinical relevance of this continuum is reinforced by redifferentiation therapy. In BRAF-driven mouse thyroid cancer, MAPK inhibition restored radioiodine incorporation, providing experimental evidence that part of the iodine-handling defect is reversible {Chakravarty, 2011}. In patients with advanced thyroid cancer, selumetinib increased radioiodine uptake in a subset of tumors, especially in RAS-mutant disease {Ho, 2013}. More recent trials using BRAF or combined BRAF/MEK inhibition have shown that targeted redifferentiation can re-enable 131I treatment in selected BRAF-mutant radioiodine-refractory tumors {Leboulleux, 2023; Tchekmedyian, 2022}. These observations create a practical need for state-based biomarkers that are not merely mutation labels.",
        "Metabolic rewiring may be one such state feature. One-carbon metabolism supports nucleotide synthesis, redox balance and methyl-group transfer, and it often becomes coupled to proliferative stress in cancer cells {Ducker, 2017}. Serine, glycine and mitochondrial folate-cycle activity provide carbon units that connect biosynthesis with epigenetic and DNA damage responses {Locasale, 2013; Amelio, 2014}. MTHFD2 is normally low in many adult tissues but recurrently induced in cancers, making it a plausible metabolic vulnerability rather than a generic housekeeping enzyme {Nilsson, 2014}. Amino-acid transporters, including SLC7A5/LAT1 and SLC1A5/ASCT2, provide upstream substrate supply and can influence mTORC1 signaling, immune context and tumor growth {Wang, 2015; Kanai, 2022; Nachef, 2021}.",
        "At the same time, dedifferentiation is not only a metabolic event. AXL signaling has been linked to EMT-like states and therapeutic resistance in multiple cancers {Antony, 2017}. In thyroid cancer, TYRO3/AXL receptor activation has direct experimental support, suggesting that RTK/EMT circuitry may intersect with thyroid cancer progression {Avilla, 2011}. Single-cell studies now add another layer by showing that papillary and anaplastic thyroid cancers contain malignant epithelial states embedded in immune, stromal and vascular ecosystems {Pu, 2021; Lu, 2023}. Therefore, a useful RAIR-like framework should be able to separate tumor epithelial signal from microenvironmental composition and should not overinterpret bulk correlations as malignant-cell specificity.",
        "In this study, we built a conservative multi-cohort framework around a composite RAIR-like score defined by one-carbon metabolism, MAPK activity, dedifferentiation/EMT and loss of thyroid differentiation. We first tested whether the score directly separated radioiodine-refractory from radioiodine-avid tumors in the available public RAI cohort. We then evaluated whether the same score mapped the PTC-PDTC-ATC dedifferentiation continuum, localized to malignant-like epithelial compartments at single-cell resolution and prioritized druggable or experimentally tractable nodes. The central aim was not to claim a ready-to-use RAI response classifier, but to identify a reproducible dedifferentiation state with translational follow-up value."
    ]),
    ("Methods", [
        ("Study design and public cohorts", [
            "This retrospective computational study integrated bulk transcriptomic, single-cell transcriptomic, clinical, mutation, protein and target-annotation resources. TCGA-THCA was used as the discovery cohort. GSE151179 was used only as an exploratory public RAI-response cohort. GSE33630 and GSE76039 were used to test the histologic dedifferentiation continuum. GSE184362 and GSE232237 were used for single-cell localization. Cohort roles and sample sizes are summarized in Table 1.",
            "TCGA-THCA RNA-seq and clinical data were accessed using TCGAbiolinks from the Genomic Data Commons, selecting primary tumor samples and STAR-count gene-expression quantification {Colaprico, 2016}. GEO expression matrices were accessed using GEOquery {Davis, 2007}. The original studies associated with GSE151179, GSE33630, GSE184362 and GSE232237 were used for clinical or histologic context and interpretation {Colombo, 2020; Hebrant, 2012; Pu, 2021; Lu, 2023}. Public de-identified data were used throughout; no new human subject recruitment or specimen generation was performed."
        ]),
        ("Bulk expression processing and module scores", [
            "For TCGA-THCA, the tpm_unstrand assay was used when available; otherwise the fallback expression assay was log2 transformed as log2(x+1). Gene identifiers were collapsed to gene symbols by retaining the feature with the largest variance when duplicate symbols were present. For GEO arrays, expression matrices were log2 transformed when required and probes were mapped to gene symbols using the relevant Bioconductor annotation packages. Duplicate probes were collapsed to the highest-variance probe per gene.",
            "The thyroid differentiation score (TDS) contained DIO1, DIO2, DUOX1, DUOX2, FOXE1, GLIS3, NKX2-1, PAX8, SLC26A4, SLC5A5, SLC5A8, TG, THRA, THRB, TPO and TSHR. The one-carbon module contained SHMT2, MTHFD2, MTHFD1L, PHGDH, PSAT1, PSPH, SLC1A5, SLC7A5, TYMS, DHFR and GART. The MAPK module contained DUSP4, DUSP5, DUSP6, ETV4, ETV5, SPRY1, SPRY2, SPRY4, CCND1 and FOSL1. The dedifferentiation/EMT module contained HMGA2, AXL, VIM, FN1, CD44, ITGA6, KRT19, EPCAM, LAMC2 and ITGB1. Iodine-handling, immune-checkpoint and CAF scores were also computed using prespecified gene sets.",
            "Within each dataset, module scores were calculated by gene-wise z transformation followed by averaging available genes in the module. The composite RAIR-like score was defined as z(One-carbon) + z(MAPK) + z(dedifferentiation/EMT) - z(TDS). This definition was fixed before downstream association testing and was applied consistently across bulk and single-cell analyses."
        ]),
        ("Clinical, mutation and pathway analyses", [
            "TCGA clinical variables were merged by patient barcode. Nodal status was coded as N1 versus N0, T category as T3/T4 versus T1/T2 and pathologic stage as III/IV versus I/II where available. Core mutation status for BRAF, NRAS, HRAS, KRAS, TERT, TP53, EIF1AX, PIK3CA and PTEN was retrieved from the cBioPortal TCGA-THCA mutation profile. RAS mutation status was defined by NRAS, HRAS or KRAS mutation.",
            "Continuous score-score associations were tested using Spearman correlation. Two-group comparisons used two-sided Wilcoxon rank-sum tests. Logistic regression estimated odds ratios per one standard deviation increase in score. Robustness models adjusted for CAF score, immune-checkpoint score, BRAF mutation and RAS mutation. Residual analyses first modeled RAIR-like or one-carbon score as a function of MAPK score, CAF score, immune-checkpoint score and driver mutation status, then tested the residual signal against iodine-handling score or nodal status.",
            "For pathway enrichment, TCGA tumors in the highest and lowest RAIR-like quartiles were compared using limma {Ritchie, 2015}. Genes used directly in the RAIR-like score and related prespecified modules were excluded before enrichment ranking to reduce self-confirmation. Gene set enrichment analysis used fgsea against MSigDB Hallmark and Reactome collections, following the standard GSEA framework {Subramanian, 2005; Liberzon, 2015}."
        ]),
        ("External validation and target prioritization", [
            "In GSE151179, the primary exploratory contrast was RAI-refractory versus RAI-avid papillary thyroid carcinoma according to the study metadata. An additional RAI-avid remission sensitivity contrast was retained for interpretation but not used as a decisive endpoint because the RAI-avid tumor group was very small. In GSE33630, ATC was compared with PTC; in GSE76039, ATC was compared with PDTC. AUCs were computed from rank statistics in the expected biological direction: higher RAIR-like and one-carbon scores but lower TDS for more dedifferentiated tumors.",
            "Candidate target genes were prioritized by TCGA association with the RAIR-like state, external validation in GSE33630/GSE76039, Human Protein Atlas thyroid cancer immunohistochemistry and ChEMBL druggability evidence. Human Protein Atlas protein and pathology resources were used to annotate normal thyroid and thyroid cancer protein expression {Uhlen, 2015; Uhlen, 2017}. ChEMBL target and activity records were used to annotate chemical tractability and mechanism-phase evidence {Gaulton, 2012; Zdrazil, 2024}."
        ]),
        ("Single-cell localization and copyKAT analysis", [
            "For GSE184362, 10X matrices were read sample by sample and normalized as log1p(counts per 10,000). For GSE232237, dense count tables were streamed for the prespecified signature and marker genes to keep the analysis tractable on the local Windows workstation; log1p counts were used for localization. Broad cell types were assigned using marker modules for epithelial/thyroid, T/NK, B/plasma, myeloid, CAF and endothelial compartments. Per-cell RAIR-like scores were computed using the same module formula after dataset-level scaling.",
            "Top-decile RAIR-like cells were summarized by broad cell type. Sample-level epithelial/thyroid RAIR-like medians were compared across tissue or disease groups and were further modeled with adjustment for myeloid, CAF, T/NK and epithelial fractions and total cell number. These models were used to evaluate whether epithelial signal was fully explained by sample composition.",
            "copyKAT v1.1.0 was used to infer aneuploid malignant-like epithelial cells from single-cell transcriptomes, following the method proposed for copy-number and clonal inference from scRNA-seq {Gao, 2021}. Because full matrices contained 282,758 cells and exceeded practical local runtime for all-cell analysis, copyKAT was run on reproducible sampled epithelial/thyroid cells together with reference immune/stromal cells. For GSE184362, an additional sensitivity run used paratumor epithelial/thyroid cells as normal reference cells through norm.cell.names. Core copyKAT parameters included id.type=\"S\", cell.line=\"no\", ngene.chr=5, min.gene.per.cell=200, LOW.DR=0.05, UP.DR=0.10, win.size=25, KS.cut=0.10, distance=\"euclidean\", genome=\"hg20\" and output.seg=\"FALSE\"."
        ]),
        ("Statistical software", [
            "Analyses were performed in R 4.5.3. Main packages included data.table 1.18.2.1, TCGAbiolinks 2.38.0, GEOquery 2.78.0, limma 3.66.0, fgsea 1.36.2, msigdbr 26.1.0, survival 3.8-6, ggplot2 4.0.3, Seurat 5.5.0, Matrix 1.7-4, copykat 1.1.0 and openxlsx 4.2.8.1. All P values were two-sided unless otherwise specified. Pathway-level enrichment used false-discovery-rate correction."
        ])
    ]),
    ("Results", [
        ("A multi-cohort framework was built around a prespecified RAIR-like dedifferentiation state", [
            "We analyzed six public resources spanning TCGA discovery, RAI-response exploration, histologic validation and single-cell localization (Table 1; Figure 1). The central exposure was a composite RAIR-like score designed to integrate one-carbon metabolism, MAPK activity and dedifferentiation/EMT while subtracting thyroid differentiation. This formulation was chosen to represent a state rather than a single pathway, because prior thyroid cancer studies have shown that iodine-handling loss, BRAF-like/MAPK signaling and high-grade dedifferentiation are related but not interchangeable {Cancer Genome Atlas Research Network, 2014; Landa, 2016}.",
            "The analysis plan deliberately treated direct RAI response and histologic dedifferentiation as different levels of evidence. GSE151179 was kept as an exploratory RAI endpoint because the RAI-avid tumor group was small and the cohort contained heterogeneous lesion contexts. GSE33630 and GSE76039 were used as the main external tests of whether the same score followed the PTC-PDTC-ATC dedifferentiation axis."
        ]),
        ("TCGA-THCA identified a MAPK-related but nonredundant RAIR-like state linked to iodine-handling loss", [
            "In 505 TCGA-THCA primary tumors, the RAIR-like score showed the expected positive correlations with one-carbon metabolism (rho=0.685), MAPK score (rho=0.812) and dedifferentiation/EMT score (rho=0.853). It correlated inversely with TDS (rho=-0.238, P=6.59e-08) and iodine-handling score (rho=-0.341, P=4.44e-15), while also showing positive associations with immune-checkpoint (rho=0.312) and CAF (rho=0.413) scores (Figure 2a).",
            "Driver context was consistent with known TCGA biology but did not fully explain the state. BRAF-mutant tumors had higher RAIR-like scores than BRAF-wild-type tumors (median 1.258 versus -0.795, P=1.27e-23), whereas RAS-mutant tumors had lower scores than RAS-wild-type tumors (median -1.645 versus 0.700, P=2.51e-08). RAIR-like score was also higher in N1 than N0 tumors (median 0.911 versus 0.032, P=1.01e-06), corresponding to an univariable OR of 1.88 per standard deviation. The T3/T4 versus T1/T2 contrast was weaker (median 0.728 versus 0.341, P=0.042), and stage III/IV was not a stable endpoint. TCGA overall survival contained only 16 events and was therefore not used as a major evidence layer.",
            "Multivariable and residual models supported a nonredundant relationship with iodine-handling loss. In a model including CAF score, immune-checkpoint score, BRAF mutation and RAS mutation, RAIR-like score remained associated with lower iodine-handling score (beta=-0.192, P=3.03e-05) and with N1 disease (OR=1.73 per standard deviation, P=3.50e-05). After removing the component explained by MAPK, CAF, immune-checkpoint and driver status, the RAIR-like residual retained a strong association with iodine-handling score (beta=-1.016, P=4.46e-37) and N1 disease (OR=2.158, P=0.00142) (Table 2; Figure 2c-d). In contrast, the one-carbon residual showed a positive association with iodine-handling score and an inverse association with N1, indicating that one-carbon metabolism alone should not be interpreted as a low-iodine surrogate.",
            "Pathway analysis added context to this result. In RAIR-like-high versus RAIR-like-low TCGA tumors, after excluding genes used directly in the score, Hallmark enrichment remained strongest for TNFA/NF-kB signaling (NES=1.96, FDR=1.99e-19), G2M checkpoint (NES=1.91, FDR=9.71e-18), interferon-gamma response (NES=1.86, FDR=2.16e-16), EMT (NES=1.76, FDR=1.25e-12) and KRAS signaling up (NES=1.71, FDR=1.65e-10). Thus, the score captured a broader inflammatory, proliferative and EMT-like transcriptional state rather than simply restating the genes used to build it (Figure 2e; Figure S3)."
        ]),
        ("Direct RAI discrimination was negative in GSE151179, whereas dedifferentiation validation was robust", [
            "The public RAI endpoint did not meet a go-forward threshold. In GSE151179, RAI-refractory tumors were not separated from RAI-avid tumors by TDS (AUC=0.579, P=0.639) or by the RAIR-like score (AUC=0.421, P=0.639) (Figure S1). The RAI-avid group contained only four tumor samples, and primary-only sensitivity analysis also remained negative. We therefore did not present the score as a direct classifier of RAI refractoriness.",
            "In contrast, the same score was reproducibly aligned with histologic dedifferentiation. In GSE33630, TDS nearly perfectly separated ATC from PTC in the expected inverse direction (AUC=0.996, P=2.33e-11), and the RAIR-like score also separated ATC from PTC (AUC=0.885, P=1.69e-05). One-carbon score showed similar ATC-versus-PTC performance (AUC=0.891, P=1.17e-05). In GSE76039, which tested a higher-grade PDTC-to-ATC contrast, TDS again showed strong inverse separation (AUC=0.988, P=1.51e-09), and RAIR-like score separated ATC from PDTC (AUC=0.871, P=4.70e-05). One-carbon score remained stable in this comparison (AUC=0.888, P=1.65e-05) (Figure 3a-b).",
            "These external results define the appropriate scope of the manuscript. The score is supported as a dedifferentiation-state marker across independent transcriptomic cohorts, but the available public RAI endpoint is insufficient to validate it as a stand-alone RAI-refractory diagnostic classifier."
        ]),
        ("Single-cell and copyKAT analyses localized RAIR-like signal to malignant-like epithelial/thyroid compartments with ecosystem contributions", [
            "Across 197,955 cells in GSE184362 and 84,803 cells in GSE232237, top-decile RAIR-like cells were enriched in epithelial/thyroid-lineage compartments but were not exclusive to them. In GSE184362, epithelial/thyroid cells accounted for 61.6% of top-decile RAIR-like cells, followed by T/NK cells (15.5%), myeloid cells (8.5%) and CAFs (6.9%). In GSE232237, epithelial/thyroid cells accounted for 48.6% of top-decile cells, with a larger myeloid fraction (26.5%) and CAF contribution (11.6%) (Figure 4a).",
            "Sample-level modeling argued against a purely compositional explanation. In GSE184362 epithelial/thyroid cells, tumor or metastatic samples had higher median RAIR-like scores than paratumor samples (1.23 versus -3.43, P=0.0039). After adjustment for myeloid, CAF, T/NK and epithelial fractions and total cell number, the tumor/metastasis indicator remained associated with epithelial RAIR-like score (beta=4.51, P=4.19e-04). In GSE232237, the composition-adjusted PTC-versus-ATC comparison was significant but in the opposite direction to a simple ATC-higher expectation (beta=2.17 for PTC relative to ATC, P=0.0216). This result was retained as a caution that broad local annotation and sample composition can influence single-cell state comparisons.",
            "copyKAT strengthened malignant-cell localization in the dataset where diploid and aneuploid epithelial populations were both available. In GSE232237, copyKAT-aneuploid epithelial/thyroid cells had substantially higher RAIR-like scores than diploid epithelial/thyroid cells (n=501 versus 52; median 2.827 versus -0.251; median difference=3.079; P=4.18e-11) (Table 2; Figure 4c-d). In GSE184362, immune/stromal references overcalled paratumor epithelial cells as aneuploid, so the stronger interpretation used paratumor epithelial cells as the normal reference. Under that reference, target epithelial cells from PTC10_T, PTC10_RightLN and PTC11_SC were almost entirely classified as aneuploid, with median RAIR-like scores of 3.44, 3.52 and 4.14, respectively. Because diploid target epithelial cells were too few in this sensitivity run, no strong within-GSE184362 aneuploid-versus-diploid comparison was claimed.",
            "Together, the single-cell data support a malignant-like epithelial/thyroid localization of the RAIR-like state, while preserving the observation that myeloid and CAF compartments contribute to the ecosystem surrounding high-score cells."
        ]),
        ("Integrated target annotation prioritized transporter, one-carbon and RTK/EMT nodes", [
            "Target annotation separated mechanistic candidates from ecology markers. SLC7A5, SLC1A5, MTHFD2 and AXL were classified as core actionable or experimentally tractable axes because they combined TCGA association, external dedifferentiation validation and some form of targetability evidence. FN1, FOSL1 and ITGB1 were consistently linked to the RAIR-like state but were interpreted mainly as ECM, adhesion or AP-1/MAPK state anchors (Table 3; Figure 5).",
            "SLC7A5 had the strongest integrated translational profile among the metabolic/transport candidates. It correlated with the RAIR-like score in TCGA (rho=0.476), was validated across both GEO dedifferentiation cohorts, showed HPA thyroid cancer high/medium staining in 75% of the small available series and had 473 ChEMBL activity records. SLC1A5 was similarly stable across transcriptomic datasets (TCGA rho=0.602) but had more limited HPA support (25% high/medium) and 143 ChEMBL activity records. MTHFD2 correlated with RAIR-like score (rho=0.596) and validated externally, but HPA thyroid cancer high/medium staining was not observed and ChEMBL evidence was sparse. This places MTHFD2 as a biologically important candidate requiring protein-level and functional validation rather than as an immediately mature drug target.",
            "AXL was strongly correlated with the RAIR-like score in TCGA (rho=0.690) and externally validated, and ChEMBL annotation supported the strongest clinical druggability signal among the candidates. However, thyroid cancer HPA immunohistochemistry did not support AXL protein detection in the small available dataset. We therefore interpreted AXL as a druggable RTK/EMT axis that requires tumor-context validation, not as a protein-validated thyroid cancer marker from public IHC alone.",
            "Within BRAF-mutant TCGA tumors, MTHFD2, SLC1A5, SLC7A5, AXL, FN1 and FOSL1 all correlated with the RAIR-like score. Only SLC7A5 and FN1 retained inverse associations with iodine-handling after false-discovery-rate correction. This reinforced the main conclusion that a composite RAIR-like state is more informative than any single metabolic gene and that the most credible translational output is a prioritized axis for experimental follow-up."
        ])
    ]),
    ("Discussion", [
        "This study supports a restrained but coherent conclusion: a composite RAIR-like state, enriched for one-carbon metabolism and amino-acid transport but anchored by MAPK activity, dedifferentiation/EMT and thyroid differentiation loss, maps the thyroid cancer dedifferentiation continuum. The result is consistent with TCGA-THCA, where BRAF-like signaling and reduced thyroid differentiation were tightly linked but still biologically separable {Cancer Genome Atlas Research Network, 2014}. It also fits the transcriptomic architecture of PDTC and ATC, in which high-grade progression involves thyroid-lineage loss together with cell-cycle, MAPK, PI3K and tumor-suppressor disruption {Landa, 2016}. The important boundary is that our public RAI-response analysis was negative in GSE151179, so the manuscript should not be read as validating a diagnostic classifier for radioiodine-refractory disease {Colombo, 2020}.",
        "The negative GSE151179 result should be treated as a boundary condition for the present study, and it is the reason the main claim is framed around dedifferentiation biology rather than direct RAI-response prediction. Public RAI-response transcriptomic cohorts remain small and uneven, and RAI avidity can be lesion-specific, treatment-specific and timing-dependent. The clinical literature already shows that preserved radioiodine uptake is prognostically meaningful, but that the transition to refractoriness is not captured by a single static variable {Durante, 2006}. Current guidelines also treat radioiodine-refractory disease as a clinical category defined by uptake, progression and therapeutic context, not simply by gene expression {Haugen, 2016}. Our data therefore support a state framework that may inform future RAI studies, rather than a ready clinical decision tool.",
        "The TCGA residual analyses are central to the paper. RAIR-like score was higher in BRAF-mutant tumors, as expected from the BRAF-like/TDS framework {Cancer Genome Atlas Research Network, 2014}. However, the RAIR-like residual remained strongly linked to iodine-handling loss and nodal disease after accounting for MAPK, immune, CAF and driver variables. This suggests that the composite state captures an integrated transcriptional phenotype downstream of, but not reducible to, the canonical BRAF/MAPK label. The observation is compatible with redifferentiation studies showing that MAPK inhibition can restore iodine uptake in experimental and clinical settings, while response remains incomplete and context-dependent {Chakravarty, 2011; Ho, 2013}. The more recent dabrafenib-trametinib and vemurafenib-based redifferentiation trials reinforce the same point: MAPK targeting can reopen the iodine-handling program in selected tumors, but does not remove the need for better state biomarkers {Leboulleux, 2023; Tchekmedyian, 2022}.",
        "One-carbon metabolism should be interpreted as part of this state, not as the state itself. Cancer cells often increase serine/glycine and folate-cycle flux to support nucleotide synthesis, redox balance and methylation reactions {Locasale, 2013; Ducker, 2017}. MTHFD2 is a particularly plausible cancer-associated enzyme because it is linked to the mitochondrial folate pathway and is frequently induced in tumors relative to adult tissues {Nilsson, 2014}. Nonetheless, our residual analysis showed that one-carbon score alone did not behave like a low-iodine or N1-risk surrogate. This is biologically plausible: one-carbon activity can rise with proliferation or biosynthetic demand without necessarily marking thyroid-lineage collapse. For this reason, MTHFD2 is best presented as a metabolic vulnerability embedded in the RAIR-like program, not as a stand-alone marker of radioiodine resistance.",
        "The transporter findings refine the translational interpretation. SLC7A5/LAT1 transports large neutral amino acids and has a well-developed cancer literature as a diagnostic and therapeutic target {Kanai, 2022}. LAT1-related amino-acid transport can influence mTORC1 signaling and growth programs, offering a mechanistic bridge between nutrient uptake and aggressive tumor states {Wang, 2015}. SLC1A5/ASCT2 is less mature clinically but has been discussed with SLC7A5/SLC3A2 as a way to reshape tumor metabolism and tumor-immune interactions {Nachef, 2021}. In our target map, SLC7A5 stood out because it combined cross-cohort transcriptomic stability, inverse iodine-handling association within BRAF-mutant tumors, HPA protein support and ChEMBL chemical evidence. That combination makes SLC7A5 a more immediately defensible translational node than MTHFD2, despite the latter's strong biological appeal.",
        "AXL provides a different type of targetability. AXL is a recognized EMT-associated receptor tyrosine kinase and resistance node across cancers {Antony, 2017}. In thyroid cancer, activation of TYRO3/AXL receptor signaling has been experimentally reported, supporting relevance beyond generic pan-cancer extrapolation {Avilla, 2011}. AXL also has a broader EMT and metastasis literature that explains why it can travel with FN1, ITGB1 and FOSL1-like state markers {Gjerdrum, 2010}. Our data support AXL as a druggable RTK/EMT axis within RAIR-like tumors, but not as a protein-validated thyroid cancer marker from HPA alone. That distinction matters because HPA staining for AXL was not supportive in the small public thyroid cancer series.",
        "The single-cell analyses were useful because they reduced a common ambiguity in bulk thyroid cancer studies: whether a high-risk state is epithelial tumor biology or simply a stromal/immune admixture. Papillary thyroid cancer single-cell work has shown that tumor progression occurs within complex epithelial and stromal ecosystems {Pu, 2021}. Anaplastic transformation studies further show that malignant epithelial programs and the microenvironment evolve together during high-grade transition {Lu, 2023}. In our data, top-decile RAIR-like cells were mostly epithelial/thyroid, and the GSE184362 composition-adjusted model remained positive after accounting for myeloid, CAF, T/NK and epithelial fractions. copyKAT added malignant-like support in GSE232237, where aneuploid epithelial cells had substantially higher RAIR-like scores than diploid epithelial cells {Gao, 2021}. These results justify the phrase malignant-like epithelial enrichment, but they do not justify saying that the signal is exclusively malignant epithelial or independent of the ecosystem.",
        "The GSE232237 sample-level direction also deserves direct discussion. The composition-adjusted PTC-versus-ATC comparison was significant in a PTC-higher direction, even though the bulk PDTC/ATC validation supported higher RAIR-like scores in more dedifferentiated tumors. This discrepancy likely reflects the limits of broad local marker annotation, sample composition, intratumoral heterogeneity and the fact that a single-cell cohort is not a balanced bulk histology validation set. We therefore used GSE232237 mainly for cell-ecology and copyKAT localization, rather than as a sample-level validation that ATC uniformly exceeds PTC in RAIR-like score.",
        "Clinically, the work sits between biomarker discovery and therapeutic prioritization. RAI-refractory differentiated thyroid cancer already has approved multikinase inhibitors, but these drugs are not state-specific and do not directly restore thyroid differentiation {Brose, 2014; Schlumberger, 2015}. ATC now has targeted options in selected molecular contexts, including BRAF/MEK inhibition for BRAF V600E disease, but outcomes remain heterogeneous and aggressive biology is often layered on top of targetable drivers {Subbiah, 2018; Subbiah, 2022; Bible, 2021}. A RAIR-like state score could eventually help select tumors for redifferentiation trials or for combination strategies that pair MAPK inhibition with metabolic, transporter or RTK/EMT targeting. That proposition remains a hypothesis until tested prospectively.",
        "Several limitations should shape the final message. First, all analyses are retrospective and based on public datasets. Second, the only direct RAI-response cohort was small and negative for the proposed score, so direct clinical prediction of RAI refractoriness is not established. Third, TCGA-THCA is dominated by differentiated papillary thyroid carcinoma, meaning that TCGA findings reflect early or partial dedifferentiation states rather than the full biology of PDTC and ATC. Fourth, single-cell annotation used broad marker-based classes and sampled copyKAT runs; full inferCNV/copyKAT across all cells and author-provided malignant labels would strengthen the claim. Fifth, HPA thyroid cancer protein evidence was based on small numbers of samples and should be treated as supportive, not definitive. Finally, MTHFD2, SLC7A5, SLC1A5 and AXL require functional validation in thyroid cancer models before being promoted as treatment targets.",
        "Despite these limitations, the study provides a usable translational framework. It separates the unsupported claim of direct RAI-response classification from the better-supported claim of a dedifferentiation-state map. It shows that the composite RAIR-like state remains informative beyond MAPK and microenvironmental covariates. It localizes much of the signal to malignant-like epithelial/thyroid compartments while retaining ecosystem context. It also turns a transcriptomic signature into a prioritized and falsifiable target map, with SLC7A5 and AXL as the more druggable near-term nodes and MTHFD2/SLC1A5 as experimentally tractable metabolic candidates."
    ]),
    ("Conclusions", [
        "A composite one-carbon metabolism-enriched RAIR-like score maps a reproducible dedifferentiation state across thyroid cancer cohorts. The state is associated with iodine-handling loss and nodal disease in TCGA, validates across independent PTC-PDTC-ATC bulk cohorts and is enriched in malignant-like epithelial/thyroid compartments at single-cell resolution. The current data do not validate direct prediction of radioiodine refractoriness, but they support a JTM-suitable translational hypothesis centered on dedifferentiation stratification and prioritization of SLC7A5, SLC1A5, MTHFD2 and AXL for experimental follow-up."
    ])
]


figure_legends = [
    ("Figure 1", "Study design and RAIR-like score definition. (a) Multi-cohort workflow from TCGA discovery to GEO validation, single-cell localization, copyKAT malignant-cell inference and target translation. (b) Modules used to define the RAIR-like score: z(One-carbon) + z(MAPK) + z(Dediff/EMT) - z(TDS). (c) Public cohorts and sample or cell counts. (d) Evidence strength assigned to the main claims."),
    ("Figure 2", "TCGA-THCA discovery analysis. (a) RAIR-like score versus iodine-handling score by driver group. (b) RAIR-like score distribution by BRAF/RAS driver context. (c) Iodine-handling score by MAPK-adjusted RAIR residual group. (d) Residual models testing RAIR-like and one-carbon components beyond MAPK, CAF, immune-checkpoint and driver status. (e) Hallmark enrichment in RAIR-like-high versus RAIR-like-low tumors after excluding signature genes."),
    ("Figure 3", "External GEO validation of the dedifferentiation continuum. (a) TDS, one-carbon and RAIR-like scores across histologic groups in GSE33630 and GSE76039. (b) AUCs for ATC versus PTC or ATC versus PDTC contrasts. (c) External validation heat map for prioritized candidate nodes."),
    ("Figure 4", "Single-cell localization and copyKAT malignant-cell inference. (a) Broad cell-type distribution among top-decile RAIR-like cells in GSE184362 and GSE232237. (b) Sample-level epithelial/thyroid RAIR-like scores in relation to microenvironmental composition. (c) copyKAT aneuploid fractions in sampled epithelial/thyroid compartments. (d) RAIR-like score by copyKAT class."),
    ("Figure 5", "Translational target prioritization. (a) Candidate target map integrating TCGA association, RAIR-high expression difference and GEO validation. (b) Human Protein Atlas thyroid cancer immunohistochemistry for selected targets. (c) ChEMBL activity and mechanism-phase evidence. (d) Integrated evidence matrix supporting transporter, one-carbon and RTK/EMT prioritization."),
    ("Figure S1", "Direct RAI endpoint in GSE151179. TDS and RAIR-like score did not distinguish RAI-refractory from RAI-avid tumors; this endpoint was treated as underpowered and negative."),
    ("Figure S2", "BRAF/MAPK residual detail. One-carbon score alone was not a low-iodine surrogate within BRAF-mutant tumors; target-gene correlations with iodine-handling were heterogeneous."),
    ("Figure S3", "Additional pathway enrichment. Reactome pathways enriched in RAIR-like-high tumors and the Hallmark gene sets with the lowest positive normalized enrichment scores; no Hallmark gene set had a negative normalized enrichment score in this contrast."),
    ("Figure S4", "CNV and copyKAT sensitivity analyses. Lightweight CNV proxy analyses and GSE184362 paratumor epithelial reference copyKAT runs define the limits of malignant-cell localization."),
    ("Figure S5", "Target annotation details. Normal thyroid HPA immunohistochemistry and detailed target evidence scoring for candidate nodes."),
]


declarations = [
    ("Ethics approval and consent to participate", "Not applicable. This study used de-identified public datasets and did not involve new human participant recruitment or new biospecimen collection."),
    ("Consent for publication", "Not applicable."),
    ("Availability of data and materials", "TCGA-THCA data are available from the Genomic Data Commons. GEO datasets used in this study are available under GSE151179, GSE33630, GSE76039, GSE184362 and GSE232237. Protein and target annotations were obtained from the Human Protein Atlas and ChEMBL. Analysis outputs and source data tables are provided in the manuscript asset folder generated with this study."),
    ("Competing interests", "The authors declare that they have no competing interests."),
    ("Funding", "This work was supported by the National Natural Science Foundation of China (82371602) and the Hunan Provincial Key R&D Program (2024JK2140)."),
    ("Acknowledgements", "Not applicable."),
]


## Revision v2 overrides based on the JTM editorial triage recommendations.
## These assignments intentionally supersede the earlier draft blocks above.
sections = [
    ("Introduction", [
        "Differentiated thyroid cancer is usually indolent, yet a clinically important subset loses radioiodine avidity or progresses despite radioiodine-based therapy. In metastatic papillary and follicular thyroid cancer, long-term outcome depends strongly on preserved iodine uptake {Durante, 2006}. Contemporary guidelines define loss of radioiodine uptake and progression after radioiodine as central features of radioiodine-refractory disease {Haugen, 2016}. Once this state is reached, systemic kinase inhibition becomes a major treatment option, but sorafenib and lenvatinib do not resolve the biologic heterogeneity of radioiodine-refractory disease {Brose, 2014; Schlumberger, 2015}.",
        "Molecular studies have placed thyroid differentiation loss at the center of aggressive thyroid cancer. TCGA-THCA separated papillary thyroid carcinoma into BRAF-like and RAS-like states and introduced thyroid differentiation score as a quantitative readout of iodine-handling and thyroid-lineage programs {Cancer Genome Atlas Research Network, 2014}. Poorly differentiated and anaplastic thyroid cancers extend this continuum through stronger cell-cycle activation, MAPK and PI3K pathway involvement, tumor-suppressor disruption and thyroid-lineage loss {Landa, 2016}. Integrated dedifferentiation analyses further support a graded transition from differentiated to high-grade thyroid cancer rather than a single discrete switch {Luo, 2021}.",
        "The therapeutic relevance of this continuum is reinforced by redifferentiation therapy. In BRAF-driven mouse thyroid cancer, MAPK inhibition restored radioiodine incorporation {Chakravarty, 2011}. In patients with advanced thyroid cancer, selumetinib increased radioiodine uptake in a subset of tumors, with a particularly strong signal in RAS-mutant disease {Ho, 2013}. More recent BRAF- or BRAF/MEK-directed redifferentiation trials have shown that 131I treatment can be re-enabled in selected BRAF-mutant radioiodine-refractory tumors {Leboulleux, 2023; Tchekmedyian, 2022}. These findings create a need for state-based frameworks that go beyond mutation labels.",
        "Metabolic rewiring is a plausible component of such a state. One-carbon metabolism supports nucleotide synthesis, redox balance and methyl-group transfer {Ducker, 2017}. Serine, glycine and mitochondrial folate-cycle activity connect biosynthesis with epigenetic and DNA damage responses {Locasale, 2013; Amelio, 2014}. MTHFD2 is frequently induced in cancer and has been proposed as a therapeutic vulnerability in the mitochondrial folate pathway {Nilsson, 2014; Ramos, 2024}. Amino-acid transporters, including SLC7A5/LAT1 and SLC1A5/ASCT2, provide substrate supply and may connect nutrient uptake with mTORC1 signaling, tumor growth and immune context {Wang, 2015; Kanai, 2022; Nachef, 2021}.",
        "Dedifferentiation also involves non-metabolic circuitry. AXL signaling is linked to EMT-like states, therapeutic resistance and metastatic behavior in multiple cancers {Antony, 2017; Gjerdrum, 2010}. In thyroid cancer, TYRO3/AXL receptor activation has experimental support, suggesting a biologically plausible RTK/EMT axis {Avilla, 2011}. Single-cell studies now show that papillary and anaplastic thyroid cancers contain malignant epithelial states embedded in immune, stromal and vascular ecosystems {Pu, 2021; Lu, 2023}. A dedifferentiation-state score therefore needs to be interpreted at both tumor-cell and ecosystem levels.",
        "Here we developed a composite dedifferentiation-state score integrating one-carbon metabolism, MAPK activity, dedifferentiation/EMT and loss of thyroid differentiation. We use RAIR-like score as a shorthand because the score contains iodine-handling loss, but the primary aim is dedifferentiation mapping rather than direct radioiodine-response prediction. We tested whether the score follows the PTC-PDTC-ATC continuum, localizes to malignant-like epithelial compartments and prioritizes translational nodes that can be tested in future thyroid cancer models and RAI-response cohorts."
    ]),
    ("Methods", [
        ("Study design and cohorts", [
            "This retrospective computational study integrated TCGA-THCA bulk RNA-seq, three GEO bulk cohorts, two single-cell RNA-seq cohorts, copyKAT-based malignant-cell inference, Human Protein Atlas immunohistochemistry and ChEMBL target annotation. TCGA-THCA was used as the discovery cohort. GSE151179 was used as an exploratory RAI-response cohort only. GSE33630 and GSE76039 tested the histologic dedifferentiation continuum. GSE184362 and GSE232237 were used for cell-ecology localization. Cohort roles and sample sizes are summarized in Table 1.",
            "TCGA-THCA RNA-seq and clinical data were accessed using TCGAbiolinks from the Genomic Data Commons, selecting primary tumor samples and STAR-count gene-expression quantification {Colaprico, 2016}. GEO expression matrices were accessed using GEOquery {Davis, 2007}. The original studies associated with GSE151179, GSE33630, GSE184362 and GSE232237 were used for clinical or histologic context {Colombo, 2020; Hebrant, 2012; Pu, 2021; Lu, 2023}. Public de-identified data were used throughout; no new human participant recruitment or biospecimen collection was performed."
        ]),
        ("Bulk expression processing and module scoring", [
            "For TCGA-THCA, the tpm_unstrand assay was used when available; otherwise the fallback expression assay was log2 transformed as log2(x+1). Gene identifiers were collapsed to gene symbols by retaining the feature with the largest variance when duplicate symbols were present. For GEO arrays, expression matrices were log2 transformed when required and probes were mapped to gene symbols using the relevant Bioconductor annotation packages. Duplicate probes were collapsed to the highest-variance probe per gene.",
            "The thyroid differentiation score contained DIO1, DIO2, DUOX1, DUOX2, FOXE1, GLIS3, NKX2-1, PAX8, SLC26A4, SLC5A5, SLC5A8, TG, THRA, THRB, TPO and TSHR. The one-carbon module contained SHMT2, MTHFD2, MTHFD1L, PHGDH, PSAT1, PSPH, SLC1A5, SLC7A5, TYMS, DHFR and GART. The MAPK module contained DUSP4, DUSP5, DUSP6, ETV4, ETV5, SPRY1, SPRY2, SPRY4, CCND1 and FOSL1. The dedifferentiation/EMT module contained HMGA2, AXL, VIM, FN1, CD44, ITGA6, KRT19, EPCAM, LAMC2 and ITGB1. Iodine-handling, immune-checkpoint and CAF scores were calculated from prespecified modules.",
            "Within each dataset, module scores were calculated by gene-wise z transformation followed by averaging available module genes. The same dataset-level z-score strategy was used for bulk and single-cell matrices, rather than Seurat::AddModuleScore, to keep the score definition identical across platforms. The composite dedifferentiation-state score was defined as z(One-carbon) + z(MAPK) + z(dedifferentiation/EMT) - z(TDS). This score is referred to as RAIR-like score only as a shorthand."
        ]),
        ("Clinical, mutation, pathway and statistical analyses", [
            "TCGA clinical variables were merged by patient barcode. Nodal status was coded as N1 versus N0, T category as T3/T4 versus T1/T2 and pathologic stage as III/IV versus I/II where available. Core mutation status for BRAF, NRAS, HRAS, KRAS, TERT, TP53, EIF1AX, PIK3CA and PTEN was retrieved from the cBioPortal TCGA-THCA mutation profile. RAS mutation status was defined by NRAS, HRAS or KRAS mutation.",
            "Continuous associations were tested using Spearman correlation. Two-group comparisons used two-sided Wilcoxon rank-sum tests. Logistic regression estimated odds ratios per one standard deviation increase in score. Robustness models adjusted for CAF score, immune-checkpoint score, BRAF mutation and RAS mutation. Residual analyses first modeled RAIR-like or one-carbon score as a function of MAPK score, CAF score, immune-checkpoint score and driver mutation status, then tested the residual against iodine-handling score or nodal status.",
            "For pathway enrichment, TCGA tumors in the highest and lowest RAIR-like quartiles were compared using limma {Ritchie, 2015}. Genes used directly in the prespecified score and related modules were excluded before enrichment ranking. Gene set enrichment analysis used fgsea against MSigDB Hallmark and Reactome collections {Subramanian, 2005; Liberzon, 2015}. FDR correction was applied within each tested gene-set collection or target-validation comparison family. Reporting followed relevant elements of TRIPOD where applicable, although this study did not develop a clinical prediction model."
        ]),
        ("RAI-response power context and external validation", [
            "In GSE151179, the exploratory contrast was RAI-refractory versus RAI-avid papillary thyroid carcinoma according to study metadata. The analyzable tumor contrast contained 35 RAI-refractory and 4 RAI-avid samples. Post hoc power context was calculated using pROC::power.roc.test. At alpha=0.05, this sample configuration provided 25.1%, 37.3%, 53.0% and 71.8% power to detect target AUCs of 0.70, 0.75, 0.80 and 0.85, respectively. Therefore, GSE151179 was treated as exploratory and not as a primary validation endpoint.",
            "In GSE33630, ATC was compared with PTC; in GSE76039, ATC was compared with PDTC. AUCs were computed from rank statistics in the expected biological direction: higher RAIR-like and one-carbon scores but lower TDS for more dedifferentiated tumors. Candidate target genes were prioritized by TCGA association with the score, external validation in GSE33630/GSE76039, HPA thyroid cancer immunohistochemistry and ChEMBL druggability evidence {Uhlen, 2015; Uhlen, 2017; Gaulton, 2012; Zdrazil, 2024}."
        ]),
        ("Single-cell localization and copyKAT analysis", [
            "For GSE184362, 10X matrices were read sample by sample and normalized as log1p(counts per 10,000). For GSE232237, dense count tables were streamed for prespecified signature and marker genes; log1p counts were used for localization. Because the public matrices were already processed and the analysis did not perform de novo clustering, no additional doublet removal, mitochondrial-percentage filtering or batch integration was applied. This choice preserved the original public cell composition but is treated as a limitation.",
            "Broad cell types were assigned using marker modules. Epithelial/thyroid markers were EPCAM, KRT8, KRT18, KRT19, PAX8, TG, TPO, SLC5A5, TSHR and SLC26A4. T/NK markers were PTPRC, CD3D, CD3E, TRAC, NKG7, GNLY and GZMB. B/plasma markers were MS4A1, CD79A, CD79B, MZB1, JCHAIN and IGHG1. Myeloid markers were LYZ, LST1, CD68, CD14, FCGR3A, S100A8 and S100A9. CAF markers were COL1A1, COL1A2, DCN, LUM, ACTA2, PDGFRA and TAGLN. Endothelial markers were PECAM1, VWF, KDR, ENG, RAMP2 and CLDN5.",
            "Top-decile RAIR-like cells were summarized by broad cell type. Sample-level epithelial/thyroid RAIR-like medians were compared across tissue or disease groups and were modeled with adjustment for myeloid, CAF, T/NK and epithelial fractions and total cell number. For GSE232237, unadjusted sample-level ATC-versus-PTC medians were also reported to make the direction of the single-cell sample-level comparison explicit.",
            "copyKAT v1.1.0 was used to infer aneuploid malignant-like epithelial cells from single-cell transcriptomes {Gao, 2021}. Because full matrices contained 282,758 cells and exceeded practical local runtime for all-cell analysis, copyKAT was run on reproducible sampled epithelial/thyroid cells together with reference immune/stromal cells. For GSE184362, a sensitivity run used paratumor epithelial/thyroid cells as normal reference cells through norm.cell.names. Core copyKAT parameters included id.type=\"S\", cell.line=\"no\", ngene.chr=5, min.gene.per.cell=200, LOW.DR=0.05, UP.DR=0.10, win.size=25, KS.cut=0.10, distance=\"euclidean\", genome=\"hg20\" and output.seg=\"FALSE\"."
        ]),
        ("Software and code availability", [
            "Analyses were performed in R 4.5.3. Main packages included data.table 1.18.2.1, TCGAbiolinks 2.38.0, GEOquery 2.78.0, limma 3.66.0, fgsea 1.36.2, msigdbr 26.1.0, survival 3.8-6, ggplot2 4.0.3, Seurat 5.5.0, Matrix 1.7-4, copykat 1.1.0 and openxlsx 4.2.8.1. All P values were two-sided unless otherwise specified. Custom scripts and processed source-data tables are available in the analysis folder generated with this study and should be deposited in a public repository before submission."
        ])
    ]),
    ("Results", [
        ("A prespecified composite dedifferentiation-state score integrates one-carbon, MAPK and EMT modules across six public cohorts", [
            "We analyzed six public resources spanning TCGA discovery, RAI-response exploration, histologic validation and single-cell localization (Table 1; Figure 1). The primary exposure was a composite dedifferentiation-state score integrating one-carbon metabolism, MAPK activity and dedifferentiation/EMT while subtracting thyroid differentiation. The shorthand RAIR-like is retained because the score contains iodine-handling loss, but direct RAI-response prediction was not the primary claim.",
            "The analysis plan treated direct RAI response and histologic dedifferentiation as different evidence layers. GSE151179 was classified as exploratory because the RAI-avid tumor group was small. GSE33630 and GSE76039 were used as the main external tests of whether the same score followed the PTC-PDTC-ATC dedifferentiation axis."
        ]),
        ("TCGA-THCA identifies a MAPK-related but nonredundant score associated with iodine-handling loss", [
            "In 505 TCGA-THCA primary tumors, the composite score correlated positively with one-carbon metabolism (rho=0.685), MAPK score (rho=0.812) and dedifferentiation/EMT score (rho=0.853). It correlated inversely with TDS (rho=-0.238, P=6.59e-08) and iodine-handling score (rho=-0.341, P=4.44e-15), while also showing positive associations with immune-checkpoint (rho=0.312) and CAF (rho=0.413) scores (Figure 2a).",
            "Driver context was consistent with known TCGA biology but did not fully explain the state. BRAF-mutant tumors had higher scores than BRAF-wild-type tumors (median 1.258 versus -0.795, P=1.27e-23), whereas RAS-mutant tumors had lower scores than RAS-wild-type tumors (median -1.645 versus 0.700, P=2.51e-08). The score was also higher in N1 than N0 tumors (median 0.911 versus 0.032, P=1.01e-06). The T3/T4 contrast was weaker (median 0.728 versus 0.341, P=0.042), and stage III/IV was not stable after adjustment.",
            "Multivariable and residual models supported a nonredundant relationship with iodine-handling loss. In a model including CAF score, immune-checkpoint score, BRAF mutation and RAS mutation, the score remained associated with lower iodine-handling score (beta=-0.192, 95% CI -0.282 to -0.102, P=3.03e-05) and N1 disease (OR=1.733, 95% CI 1.336 to 2.249, P=3.50e-05). After removing the component explained by MAPK, CAF, immune-checkpoint and driver status, the residual retained a strong association with iodine-handling score (beta=-1.016, 95% CI -1.160 to -0.872, P=4.46e-37) and N1 disease (OR=2.158, 95% CI 1.345 to 3.461, P=0.001) (Table 2; Figure 2c-d). One-carbon residual signal alone did not behave as a low-iodine surrogate.",
            "In RAIR-like-high versus RAIR-like-low TCGA tumors, after excluding signature genes, Hallmark enrichment remained strongest for TNFA/NF-kB signaling (NES=1.96, FDR=1.99e-19), G2M checkpoint (NES=1.91, FDR=9.71e-18), interferon-gamma response (NES=1.86, FDR=2.16e-16), EMT (NES=1.76, FDR=1.25e-12) and KRAS signaling up (NES=1.71, FDR=1.65e-10). These pathways indicate that the score tracks inflammatory, proliferative and EMT-like programs beyond the genes used to define the score (Figure 2e; Figure S3)."
        ]),
        ("GSE151179 is underpowered for RAI-response classification, whereas independent cohorts validate dedifferentiation mapping", [
            "As a prespecified exploratory analysis, we tested whether the score discriminated RAI-refractory from RAI-avid tumors in GSE151179. The analyzable contrast contained 35 RAI-refractory and 4 RAI-avid tumors. Power was limited for moderate effects: 25.1% for AUC=0.70, 37.3% for AUC=0.75 and 53.0% for AUC=0.80 at alpha=0.05. Neither TDS (AUC=0.579, P=0.639) nor the composite score (AUC=0.421, P=0.639) separated the two groups (Figure S1). This result was interpreted as an underpowered exploratory endpoint, not as evidence that the dedifferentiation score lacks RAI-response relevance.",
            "The same score was reproducibly aligned with histologic dedifferentiation. In GSE33630, TDS separated ATC from PTC in the expected inverse direction (AUC=0.996, P=2.33e-11), and the composite score also separated ATC from PTC (AUC=0.885, P=1.69e-05). One-carbon score showed similar ATC-versus-PTC performance (AUC=0.891, P=1.17e-05). In GSE76039, TDS separated ATC from PDTC (AUC=0.988, P=1.51e-09), and the composite score separated ATC from PDTC (AUC=0.871, P=4.70e-05). One-carbon score remained stable in this high-grade contrast (AUC=0.888, P=1.65e-05) (Figure 3a-b)."
        ]),
        ("Single-cell and copyKAT analyses support malignant-like epithelial localization but not sample-level ATC validation", [
            "Across 197,955 cells in GSE184362 and 84,803 cells in GSE232237, top-decile high-score cells were enriched in epithelial/thyroid compartments but were not exclusive to them. In GSE184362, epithelial/thyroid cells accounted for 61.6% of top-decile cells, followed by T/NK cells (15.5%), myeloid cells (8.5%) and CAFs (6.9%). In GSE232237, epithelial/thyroid cells accounted for 48.6% of top-decile cells, with larger myeloid (26.5%) and CAF (11.6%) contributions (Figure 4a).",
            "GSE184362 provided the clearest sample-level epithelial signal. In epithelial/thyroid cells, tumor or metastatic samples had higher median scores than paratumor samples (1.23 versus -3.43, P=0.0039). After adjustment for myeloid, CAF, T/NK and epithelial fractions and total cell number, the tumor/metastasis indicator remained associated with epithelial score (beta=4.51, 95% CI 2.57 to 6.46, P=4.19e-04).",
            "In GSE232237, sample-level histology comparison was limited by small sample number and overlapping distributions. The unadjusted epithelial/thyroid median was 1.94 in ATC samples (n=4) and 2.69 in PTC samples (n=7), with no significant difference (P=0.395). After composition adjustment, the PTC-versus-ATC coefficient was positive (beta=2.17, 95% CI 0.52 to 3.81, P=0.0216). This cohort was therefore used for cell-type localization and copyKAT classification rather than as a sample-level validation that ATC exceeds PTC.",
            "copyKAT strengthened malignant-cell localization within GSE232237. copyKAT-aneuploid epithelial/thyroid cells had higher scores than diploid epithelial/thyroid cells (n=501 versus 52; median 2.827 versus -0.251; median difference=3.079; P=4.18e-11) (Table 2; Figure 4c-d). In GSE184362, immune/stromal references overcalled paratumor epithelial cells as aneuploid, so the stricter interpretation used paratumor epithelial cells as the normal reference. Under that reference, target epithelial cells from PTC10_T, PTC10_RightLN and PTC11_SC were almost entirely classified as aneuploid, with median scores of 3.44, 3.52 and 4.14, respectively. Diploid target epithelial cells were too few for a strong within-GSE184362 aneuploid-versus-diploid comparison."
        ]),
        ("Integrated target annotation prioritizes transporter, one-carbon and RTK/EMT nodes", [
            "Target annotation separated mechanistic candidates from ecology markers. SLC7A5, SLC1A5, MTHFD2 and AXL were classified as core translational axes because they combined TCGA association, external dedifferentiation validation and some targetability evidence. FN1, FOSL1 and ITGB1 were consistently linked to the score but were interpreted mainly as ECM, adhesion or AP-1/MAPK state anchors (Table 3; Figure 5).",
            "SLC7A5 had the strongest integrated transporter profile. It correlated with the score in TCGA (rho=0.476), was validated across both GEO dedifferentiation cohorts, showed limited but directionally supportive HPA thyroid cancer high/medium staining in 3/4 public cases and had 473 ChEMBL activity records. SLC1A5 was similarly stable across transcriptomic datasets (TCGA rho=0.602) but had weaker HPA support (1/4 high/medium) and 143 ChEMBL activity records.",
            "MTHFD2 correlated with the score (rho=0.596) and validated externally, but HPA thyroid cancer high/medium staining was not observed (0/2) and ChEMBL evidence was sparse. It is therefore a biologically important metabolic candidate requiring protein-level and functional validation. AXL correlated strongly with the score (rho=0.690) and had the strongest ChEMBL actionability signal among candidates, but HPA thyroid cancer immunohistochemistry did not detect AXL in the small public series (0/4). AXL was therefore interpreted as a druggable RTK/EMT axis requiring thyroid tumor-context validation, not as a protein-validated marker from public IHC alone.",
            "Within BRAF-mutant TCGA tumors, MTHFD2, SLC1A5, SLC7A5, AXL, FN1 and FOSL1 all correlated with the composite score. Only SLC7A5 and FN1 retained inverse associations with iodine-handling after false-discovery-rate correction. This result further supports prioritizing a composite state and a target axis, rather than elevating any single metabolic gene to a stand-alone RAI-refractory marker."
        ])
    ]),
    ("Discussion", [
        ("Summary of main findings", [
            "This study maps a one-carbon-metabolism-enriched composite dedifferentiation-state score across thyroid cancer cohorts. The score tracked iodine-handling loss and nodal disease in TCGA-THCA, validated across independent PTC-PDTC-ATC bulk cohorts, localized mainly to epithelial/thyroid compartments at single-cell level and prioritized SLC7A5, SLC1A5, MTHFD2 and AXL as translational nodes. The direct public RAI-response cohort was underpowered and did not support clinical RAI-response classification.",
            "This framing is important. The study does not validate a diagnostic classifier for radioiodine-refractory disease. Instead, it provides a reproducible dedifferentiation atlas and a target-prioritization framework that can be applied to better-powered RAI-response cohorts."
        ]),
        ("Relationship to prior dedifferentiation literature", [
            "The results are consistent with TCGA-THCA, where BRAF-like signaling and reduced thyroid differentiation were related but biologically separable {Cancer Genome Atlas Research Network, 2014}. They also align with PDTC and ATC transcriptomic studies showing that high-grade progression involves thyroid-lineage loss together with cell-cycle, MAPK, PI3K and tumor-suppressor programs {Landa, 2016}. Compared with prior integrated dedifferentiation work {Luo, 2021}, the present analysis adds a focused one-carbon and amino-acid transport axis, single-cell localization, copyKAT malignant-cell inference and target-annotation layers.",
            "The GSE151179 result should be interpreted as a limitation of available public RAI-response resources. The cohort contained only four RAI-avid tumors in the analyzable contrast, and power was limited for moderate AUC effects. Clinical literature already shows that preserved radioiodine uptake is prognostically meaningful, but the transition to refractoriness is lesion-specific and treatment-context dependent {Durante, 2006; Haugen, 2016}. Future RAI-response cohorts should therefore be designed prospectively with adequate sample size and lesion-level annotation."
        ]),
        ("Implications of the TCGA residual analysis", [
            "The TCGA residual analyses are central. The composite score was higher in BRAF-mutant tumors, as expected from the BRAF-like/TDS framework {Cancer Genome Atlas Research Network, 2014}. However, the residual score remained strongly linked to iodine-handling loss and nodal disease after accounting for MAPK, immune, CAF and driver variables. The score therefore captures an integrated transcriptional phenotype downstream of, but not reducible to, the canonical BRAF/MAPK label.",
            "This observation is compatible with redifferentiation studies showing that MAPK inhibition can restore iodine uptake in experimental and clinical settings, while response remains incomplete and context dependent {Chakravarty, 2011; Ho, 2013}. Recent dabrafenib-trametinib and vemurafenib-based redifferentiation trials reinforce the same point: MAPK targeting can reopen the iodine-handling program in selected tumors, but mutation labels alone do not define the full dedifferentiation state {Leboulleux, 2023; Tchekmedyian, 2022}.",
            "One-carbon metabolism should be interpreted as a component of the state, not as the state itself. Cancer cells often increase serine/glycine and folate-cycle flux to support nucleotide synthesis, redox balance and methylation reactions {Locasale, 2013; Ducker, 2017}. MTHFD2 is a plausible cancer-associated enzyme in this pathway {Nilsson, 2014; Ramos, 2024}. In our data, however, one-carbon score alone did not behave as a low-iodine or N1-risk surrogate. MTHFD2 is best presented as an experimentally testable metabolic candidate embedded within the composite program."
        ]),
        ("Single-cell localization and the GSE232237 inconsistency", [
            "The single-cell analyses address whether the bulk signal reflects tumor epithelial biology or stromal/immune admixture. Papillary thyroid cancer single-cell work has shown that progression occurs within complex epithelial and stromal ecosystems {Pu, 2021}. Anaplastic transformation studies further show that malignant epithelial programs and the microenvironment evolve together during high-grade transition {Lu, 2023}. In our data, top-decile cells were mostly epithelial/thyroid, and the GSE184362 composition-adjusted model remained positive after accounting for myeloid, CAF, T/NK and epithelial fractions.",
            "GSE232237 produced the main single-cell inconsistency. The unadjusted ATC-versus-PTC sample-level comparison was not significant, and the composition-adjusted model had a PTC-higher direction. This discrepancy is most consistent with three non-mutually exclusive explanations: the cohort contained only four ATC and seven PTC samples for sample-level comparison; ATC samples had heterogeneous epithelial fractions; and single-cell distributions may not be well represented by sample medians. Critically, the copyKAT-aneuploid versus diploid epithelial comparison within GSE232237 was internally consistent with the dedifferentiation framework. This supports malignant-cell association even where sample-level histology comparison is inconclusive."
        ]),
        ("Translational implications", [
            "The study provides three translational outputs. First, the composite score is a quantitative stratification framework that can be applied to future thyroid cancer transcriptomic cohorts, including retrospective RAI-response series and prospective biomarker-discovery studies. Second, the target map is hypothesis-generating rather than definitive. SLC7A5/LAT1 has existing inhibitor experience in solid tumors, including JPH203 biomarker analyses {Okano, 2020}. AXL has clinical-stage inhibitor evidence outside thyroid cancer, including bemcentinib-based trials {Bhalla, 2023}. MTHFD2 has selective tool-compound evidence, including DS18561882 {Kawai, 2019}. These examples make the prioritized nodes experimentally tractable, but they do not establish thyroid cancer efficacy.",
            "Third, the score may help design redifferentiation studies. Because the composite residual retained association with iodine-handling loss after MAPK adjustment, future trials could stratify tumors into MAPK-dominant and MAPK-independent dedifferentiation groups. This would test whether baseline state scores predict 131I uptake recovery independently of BRAF/RAS genotype.",
            "The framework generates concrete predictions. SLC7A5 inhibition should have greater functional impact in high-score patient-derived thyroid cancer models than in low-score models. Baseline score should correlate with 131I uptake recovery in redifferentiation trials independently of driver genotype. In matched RAI-avid and RAI-refractory longitudinal samples, the score should increase during the transition to refractoriness."
        ]),
        ("Limitations and future work", [
            "Several limitations remain. First, all analyses are retrospective and based on public datasets. Second, GSE151179 was too small on the RAI-avid side to support direct RAI-response biomarker validation. Third, TCGA-THCA is dominated by differentiated papillary thyroid carcinoma, so TCGA findings reflect early or partial dedifferentiation rather than the full PDTC/ATC spectrum.",
            "Fourth, single-cell annotation used broad marker-based classes without additional doublet removal, mitochondrial filtering or batch integration. Fifth, copyKAT was run on reproducible sampled cell sets rather than all cells, and author-curated malignant labels were not available in the local analysis. Sixth, HPA thyroid cancer immunohistochemistry involved very small denominators and was used only as supportive annotation. Seventh, SLC7A5, SLC1A5, MTHFD2 and AXL require functional validation before being promoted as therapeutic targets in thyroid cancer."
        ])
    ]),
    ("Conclusions", [
        "A one-carbon-metabolism-enriched composite dedifferentiation-state score maps thyroid cancer dedifferentiation across bulk and single-cell cohorts. It is associated with iodine-handling loss and nodal disease in TCGA, validates across independent PTC-PDTC-ATC bulk cohorts and is enriched in malignant-like epithelial/thyroid compartments. The current data do not validate direct prediction of radioiodine refractoriness. They support a clinically relevant, hypothesis-generating framework for dedifferentiation stratification and for prioritizing SLC7A5, SLC1A5, MTHFD2 and AXL for functional validation."
    ])
]

figure_legends = [
    ("Figure 1", "Study design and composite dedifferentiation-state score definition. (a) Multi-cohort workflow from TCGA discovery to GEO validation, single-cell localization, copyKAT malignant-cell inference and target translation. (b) Score modules: z(One-carbon) + z(MAPK) + z(Dediff/EMT) - z(TDS). (c) Public cohorts and sample or cell counts. (d) Evidence strength assigned to the main claims."),
    ("Figure 2", "TCGA-THCA discovery analysis. (a) Composite score versus iodine-handling score by driver group. (b) Score distribution by BRAF/RAS driver context. (c) Iodine-handling score by MAPK-adjusted residual group. (d) Residual models beyond MAPK, CAF, immune-checkpoint and driver status. (e) Hallmark enrichment in high-score versus low-score tumors after excluding signature genes."),
    ("Figure 3", "External validation of the dedifferentiation continuum. (a) TDS, one-carbon and composite scores across histologic groups in GSE33630 and GSE76039. (b) AUCs for ATC versus PTC or ATC versus PDTC contrasts. (c) External validation heat map for prioritized candidate nodes."),
    ("Figure 4", "Single-cell localization and copyKAT malignant-cell inference. (a) Broad cell-type distribution among top-decile high-score cells in GSE184362 and GSE232237. (b) Sample-level epithelial/thyroid score in relation to microenvironmental composition; GSE232237 is interpreted as localization rather than sample-level ATC validation. (c) copyKAT aneuploid fractions in sampled epithelial/thyroid compartments. (d) Score by copyKAT class."),
    ("Figure 5", "Translational target prioritization. (a) Candidate target map integrating TCGA association, expression difference and GEO validation. (b) Human Protein Atlas thyroid cancer immunohistochemistry counts for selected targets; denominators are small and used as supportive annotation. (c) ChEMBL activity and mechanism-phase evidence. (d) Integrated evidence matrix supporting transporter, one-carbon and RTK/EMT prioritization."),
    ("Figure S1", "Exploratory RAI endpoint and power context in GSE151179. TDS and the composite score did not distinguish RAI-refractory from RAI-avid tumors in a contrast with 35 RAI-refractory and 4 RAI-avid tumors."),
    ("Figure S2", "BRAF/MAPK residual detail. One-carbon score alone was not a low-iodine surrogate within BRAF-mutant tumors; target-gene correlations with iodine-handling were heterogeneous."),
    ("Figure S3", "Additional pathway enrichment. Reactome pathways enriched in high-score tumors and the Hallmark gene sets with the lowest positive normalized enrichment scores; no Hallmark gene set had a negative normalized enrichment score in this contrast."),
    ("Figure S4", "CNV and copyKAT sensitivity analyses. Lightweight CNV proxy analyses and GSE184362 paratumor epithelial reference copyKAT runs define localization limits."),
    ("Figure S5", "Target annotation details. Normal thyroid HPA immunohistochemistry and detailed target evidence scoring for candidate nodes."),
]

declarations = [
    ("Ethics approval and consent to participate", "Not applicable. This study used de-identified public datasets and did not involve new human participant recruitment or new biospecimen collection."),
    ("Consent for publication", "Not applicable."),
    ("Availability of data and materials", "TCGA-THCA data are available from the Genomic Data Commons. GEO datasets used in this study are available under GSE151179, GSE33630, GSE76039, GSE184362 and GSE232237. Protein and target annotations were obtained from the Human Protein Atlas and ChEMBL. Custom scripts and processed source-data tables are currently available in the local analysis folder and should be deposited in a public repository with a persistent identifier before submission."),
    ("Competing interests", "The authors declare that they have no competing interests."),
    ("Funding", "This work was supported by the National Natural Science Foundation of China (82371602) and the Hunan Provincial Key R&D Program (2024JK2140). The funding sources had no role in study design, data analysis, interpretation or manuscript preparation."),
    ("Authors' contributions", "Shi Yin and Chang Shi conceived the study. Shi Yin performed the computational analyses and drafted the manuscript. Li Shuishi contributed clinical interpretation and manuscript revision. Chang Shi supervised the study and revised the manuscript. All authors reviewed and approved the final manuscript."),
    ("Acknowledgements", "Not applicable."),
]


def set_cell_text(cell, text, size=8, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in kwargs[edge].items():
                element.set(qn("w:{}".format(key)), str(value))


def set_cell_width(cell, width_in):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def table_three_line(doc, caption, headers, rows, widths, font_size=8):
    cap = doc.add_paragraph()
    cap.paragraph_format.keep_with_next = True
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(3)
    r = cap.add_run(caption)
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(10)

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.autofit = False
    tbl.allow_autofit = False

    tbl_pr = tbl._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    header_cells = tbl.rows[0].cells
    tr_pr = tbl.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for i, h in enumerate(headers):
        set_cell_text(header_cells[i], h, size=font_size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_width(header_cells[i], widths[i])
        set_cell_shading(header_cells[i], "F4F6F9")
        set_cell_border(
            header_cells[i],
            top={"val": "single", "sz": "12", "space": "0", "color": "000000"},
            bottom={"val": "single", "sz": "8", "space": "0", "color": "000000"},
            left={"val": "nil"},
            right={"val": "nil"},
        )

    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 2, 3, 4, 5) and len(str(val)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], val, size=font_size, bold=False, align=align)
            set_cell_width(cells[i], widths[i])
            set_cell_border(
                cells[i],
                top={"val": "nil"},
                bottom={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
            )

    bottom_cells = tbl.rows[-1].cells
    for c in bottom_cells:
        set_cell_border(
            c,
            bottom={"val": "single", "sz": "12", "space": "0", "color": "000000"},
            left={"val": "nil"},
            right={"val": "nil"},
        )
    doc.add_paragraph()
    return tbl


def add_para(doc, text, style=None, bold_label=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(12)
        r2 = p.add_run(" " + text)
    else:
        r2 = p.add_run(text)
    r2.font.name = "Times New Roman"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r2.font.size = Pt(12)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(14 if level == 1 else 12)
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(8)
    for name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
    return doc


def build_manuscript():
    doc = setup_doc()
    doc.core_properties.author = "Shi Yin; Li Shuishi; Chang Shi"
    doc.core_properties.title = title
    doc.core_properties.subject = "Journal of Translational Medicine manuscript draft"
    doc.core_properties.keywords = keywords

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(authors)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(11)

    for line in affiliations:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(9)

    add_heading(doc, "Abstract", 1)
    for k, v in abstract_sections.items():
        add_para(doc, v, bold_label=f"{k}:")
    add_para(doc, keywords, bold_label="Keywords:")

    for sec_title, content in sections:
        add_heading(doc, sec_title, 1)
        for item in content:
            if isinstance(item, tuple):
                add_heading(doc, item[0], 2)
                for para in item[1]:
                    add_para(doc, para)
            else:
                add_para(doc, item)

    add_heading(doc, "Declarations", 1)
    for label, text in declarations:
        add_para(doc, text, bold_label=f"{label}:")

    add_heading(doc, "References", 1)
    for i, (author, year, ref_title, journal, doi, pmid) in enumerate(references, 1):
        doi_part = f" doi:{doi}." if doi else ""
        add_para(doc, f"{i}. {author}. {ref_title} {journal}. {year}.{doi_part} PMID:{pmid}.")

    add_heading(doc, "Tables", 1)
    table_three_line(
        doc,
        "Table 1. Public cohorts used in this study.",
        ["Cohort", "Modality", "Role", "N"],
        table1,
        [1.25, 1.15, 3.2, 0.75],
        font_size=8.5,
    )
    table_three_line(
        doc,
        "Table 2. Key evidence supporting the composite dedifferentiation-state score.",
        ["Evidence", "Cohort", "Effect type", "Estimate (95% CI)", "P value", "Interpretation"],
        table2,
        [1.45, 0.78, 0.72, 1.05, 0.65, 1.85],
        font_size=7.2,
    )
    table_three_line(
        doc,
        "Table 3. Prioritized target and state-marker map.",
        ["Gene", "Tier", "Axis", "TCGA rho", "GEO", "HPA high/medium n/N", "ChEMBL", "Interpretation"],
        table3,
        [0.43, 0.72, 0.75, 0.52, 0.33, 0.72, 1.03, 2.0],
        font_size=6.7,
    )

    add_heading(doc, "Figure legends", 1)
    for label, text in figure_legends:
        add_para(doc, text, bold_label=f"{label}.")

    doc.save(MANUSCRIPT_DOCX)
    return doc


def build_notes():
    doc = setup_doc()
    doc.core_properties.author = "Shi Yin; Li Shuishi; Chang Shi"
    doc.core_properties.title = "Revision notes for RAIR-like thyroid cancer manuscript"
    doc.core_properties.subject = "Manuscript revision and follow-up notes"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("正文自审与后续补充说明")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(16)

    add_heading(doc, "本轮已完成的核查和修改", 1)
    notes = [
        "已将主线从“直接预测 RAI 难治”调整为“RAIR-like 去分化状态及其转化靶点图谱”，避免与 GSE151179 阴性结果冲突。",
        "结果部分按证据强度重排：TCGA 发现、GEO 去分化验证、GSE151179 阴性探索、单细胞/ copyKAT 定位、靶点图谱。",
        "讨论中逐句区分事实、解释和假设；未把 MTHFD2 或 one-carbon 模块单独写成 RAI 难治驱动因素。",
        "AXL 被写为可药 RTK/EMT 轴，但同时说明 HPA thyroid cancer IHC 不支持其蛋白表达，避免过度承诺。",
        "GSE232237 单细胞 PTC 与 ATC 方向不符合简单预期，正文已作为异质性和方法限制处理，而不是强行作为阳性验证。",
        "方法部分写成预设评分和分析流程，避免出现根据结果倒推方法的语言。",
        "表格已压缩为三线表形式，Table 3 由原始大表改为可读性更高的靶点优先级表。",
    ]
    for n in notes:
        add_para(doc, n)

    add_heading(doc, "投稿前建议补充或确认", 1)
    todo = [
        "确认作者排序、通讯作者设置和作者贡献。PDF 中通讯作者信息包含 Chang Shi，同时 ORCID 栏列出 Shi Yin 和 Chang Shi；当前稿件暂按 Shi Yin 与 Chang Shi 共同通讯处理。",
        "确认利益冲突声明、伦理声明和数据可用性声明是否符合课题组真实情况。",
        "若时间允许，优先补充至少一个实验或小型临床验证：SLC7A5/MTHFD2/AXL 的 IHC、qPCR、Western blot 或细胞模型干预。",
        "若后续可获得作者原始 single-cell malignant labels，可替换 broad marker annotation，并将 copyKAT 作为补充验证。",
        "若能获得更大的 RAI-refractory 队列，应作为正式临床验证；当前 GSE151179 不足以支撑直接 RAI 分类。",
        "投稿前建议让共同作者逐条核对所有基金号、单位英文、邮箱和 ORCID。",
        "图文件已在 manuscript_assets/figures 中生成；正式投稿时需要按 JTM/BMC 系统分别上传 TIFF/PDF 或合并图。",
    ]
    for n in todo:
        add_para(doc, n)

    add_heading(doc, "主要审稿风险", 1)
    risks = [
        "缺少湿实验验证仍是最大风险。JTM 可以接受高质量公共数据整合，但转化医学期刊通常希望看到蛋白或功能层面的外部支撑。",
        "RAI endpoint 阴性可能被审稿人质疑题目中的 RAIR-like 表述。当前正文已明确 RAIR-like 是状态名，不是已验证的 RAI 难治诊断器。",
        "copyKAT 是抽样运行，不是全量 28 万细胞运行；正文已公开这一限制。",
        "HPA thyroid cancer IHC 样本量小，不能作为强蛋白验证。正文仅把它作为支持性转化注释。",
    ]
    for n in risks:
        add_para(doc, n)

    doc.save(NOTES_DOCX)


def write_markdown():
    lines = [f"# {title}", "", authors, ""]
    lines.extend(affiliations)
    lines += ["", "## Abstract"]
    for k, v in abstract_sections.items():
        lines.append(f"**{k}:** {v}")
    lines += ["", f"**Keywords:** {keywords}", ""]
    for sec_title, content in sections:
        lines += [f"## {sec_title}", ""]
        for item in content:
            if isinstance(item, tuple):
                lines += [f"### {item[0]}", ""]
                for para in item[1]:
                    lines += [para, ""]
            else:
                lines += [item, ""]
    lines += ["## Declarations", ""]
    for label, text in declarations:
        lines += [f"**{label}:** {text}", ""]
    lines += ["## References", ""]
    for i, (author, year, ref_title, journal, doi, pmid) in enumerate(references, 1):
        doi_part = f" doi:{doi}." if doi else ""
        lines.append(f"{i}. {author}. {ref_title} {journal}. {year}.{doi_part} PMID:{pmid}.")
    lines += ["", "## Tables", "", "Tables are included in the DOCX as formatted three-line tables.", "", "## Figure legends", ""]
    for label, text in figure_legends:
        lines += [f"**{label}.** {text}", ""]
    MANUSCRIPT_MD.write_text("\n".join(lines), encoding="utf-8")


def build_notes():
    doc = setup_doc()
    doc.core_properties.author = "Shi Yin; Li Shuishi; Chang Shi"
    doc.core_properties.title = "Revision notes for JTM editorial triage recommendations"
    doc.core_properties.subject = "Revision checklist and remaining author actions"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Revision notes for the JTM editorial triage recommendations")
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(16)

    add_heading(doc, "Completed Revisions", 1)
    completed = [
        "Reframed the manuscript from a RAIR-like classifier narrative to a dedifferentiation-state atlas with prioritized translational nodes.",
        "Changed the title and abstract so that RAIR-like is a shorthand for the score, not the central unsupported claim.",
        "Added GSE151179 power context using pROC::power.roc.test. The true analyzable contrast was 35 RAI-refractory versus 4 RAI-avid tumors; power was 25.1%, 37.3%, 53.0% and 71.8% for target AUCs of 0.70, 0.75, 0.80 and 0.85.",
        "Added GSE232237 unadjusted sample-level epithelial RAIR comparison: ATC n=4, PTC n=7, median 1.94 versus 2.69, P=0.395. The cohort is now framed as localization and copyKAT support, not sample-level ATC validation.",
        "Added explicit marker lists for broad single-cell annotation and clarified that no additional doublet removal, mitochondrial filtering or batch integration was performed on the public matrices.",
        "Added a Discussion subsection on translational implications and testable predictions.",
        "Split limitations into a dedicated subsection and made copyKAT sampling, HPA denominators and lack of wet-lab validation explicit.",
        "Revised Table 1 numbers with thousands separators, Table 2 effect types and 95% CIs where available, and Table 3 HPA high/medium counts as n/N rather than percentages.",
        "Added CRediT-style author contribution language and clarified that funding sources had no role in the study.",
    ]
    for item in completed:
        add_para(doc, item)

    add_heading(doc, "New Analysis Outputs", 1)
    outputs = [
        "analysis/scripts/17_revision_sensitivity_analyses.R",
        "analysis/results/revision_sensitivity/GSE151179_RAI_AUC_power_context.csv",
        "analysis/results/revision_sensitivity/GSE232237_unadjusted_epithelial_RAIR_test.csv",
        "analysis/results/revision_sensitivity/GSE232237_sample_level_epithelial_RAIR_summary.csv",
        "analysis/results/revision_sensitivity/revision_sensitivity_summary.txt",
    ]
    for item in outputs:
        add_para(doc, item)

    add_heading(doc, "Items Still Requiring Author Input Before Submission", 1)
    todos = [
        "Create a public GitHub repository and Zenodo DOI for the code and processed source-data tables, then replace the current code-availability wording.",
        "Decide whether to keep the current author order and the current co-corresponding author wording for Shi Yin and Chang Shi.",
        "Confirm the CRediT author contribution statement with all authors.",
        "Consider adding wet-lab or clinical validation if feasible. The highest-impact options are SLC7A5/MTHFD2/AXL IHC in a thyroid cancer tissue series or functional testing in thyroid cancer models.",
        "If author-provided malignant labels for GSE232237/GSE184362 become available, repeat the malignant-cell localization analysis using those labels as a sensitivity analysis.",
        "Convert in-text citations to the final JTM/Vancouver style during the final reference-formatting stage, as requested by the user.",
        "Consider redesigning Figure 1 into a graphical abstract and further separating Figure 5 from Figure S5 before final submission.",
    ]
    for item in todos:
        add_para(doc, item)

    add_heading(doc, "Residual Editorial Risks", 1)
    risks = [
        "The manuscript remains entirely computational. This is now handled honestly as a hypothesis-generating translational framework, but wet-lab validation would still improve JTM acceptance probability.",
        "GSE151179 cannot validate direct RAI-response prediction because the RAI-avid group is too small. The manuscript now states this clearly.",
        "GSE232237 has a sample-level direction that does not validate ATC > PTC. The revised text now separates sample-level histology validation from copyKAT-based malignant epithelial localization.",
        "HPA thyroid cancer immunohistochemistry has very small denominators. The revised tables now show n/N and the discussion treats HPA as supportive only.",
    ]
    for item in risks:
        add_para(doc, item)

    doc.save(NOTES_DOCX)


def scan_docx_for_forbidden(path):
    terms = ["TODO", "password", "123456", "copy this", "AI-generated"]
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml")
    root = ET.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    text = " ".join(t.text or "" for t in root.findall(".//w:t", ns))
    found = [t for t in terms if t.lower() in text.lower()]
    return found


if __name__ == "__main__":
    build_manuscript()
    build_notes()
    write_markdown()
    bad = scan_docx_for_forbidden(MANUSCRIPT_DOCX)
    if bad:
        raise SystemExit(f"Forbidden terms in manuscript: {bad}")
    print(MANUSCRIPT_DOCX)
    print(NOTES_DOCX)
    print(MANUSCRIPT_MD)
